import os
import re

import docx
from io import StringIO
from parse_xml import is_number

from docx.shared import Inches, Pt
from win32com.client import Dispatch
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


class ExportReport:
    def __init__(self, module_path, data_path, output_path, max_path, prr_path, ppo_path, prc_path):
        self.module_path = module_path
        self.data_path = data_path
        self.output_path = output_path
        self.list_equipment_nodes = []

        self.max_path = max_path
        self.prr_path = prr_path
        self.ppo_path = ppo_path
        self.prc_path = prc_path


    @staticmethod
    def analysis_line(line):
        """
        用于解析除开110,111的其他case
        :param line:
        :return:
        """
        data_list = line.split()
        new_data = list(reversed(data_list))
        a = ''
        for i, j in enumerate(data_list[-1::-1]):
            if is_number(j):
                a = i
                break
        return data_list[2], new_data[a + 3], new_data[a + 1], new_data[a]


    @staticmethod
    def analysis_line2(line):
        """
        解析case110, case111
        :param line:
        :return:
        """
        data_list = line.split()
        new_data = list(reversed(data_list))
        a = ''
        for i, j in enumerate(data_list[-1::-1]):
            if is_number(j):
                a = i
                break
        return data_list[2], new_data[a + 2], new_data[a + 1], new_data[a]

    def open_docx(self, word_path):
        """
        将doc的文档格式，转为docx后缀后才能读取
        :return:  返回转换好的文件路径
        """
        word = Dispatch('Word.Application')
        word.Visible = True
        word.DisplayAlerts = True
        doc = word.Documents.Open(word_path)

    # def to_docx(self, word_path):
    #     """
    #     将doc的文档格式，转为docx后缀后才能读取
    #     :return:  返回转换好的文件路径
    #     """
    #     word = Dispatch('WPS.Application')
    #     word.Visible = True
    #     word.DisplayAlerts = True
    #     doc = word.Documents.Open(word_path)
    #
    #     doc.SaveAs('{}x'.format(self.module_path), 12)
    #     new_module_path = '{}x'.format(self.module_path)
    #     doc.Close()
    #     word.Quit()
    #     print('success')
    #     return new_module_path

    def read_template(self):
        """
        读取word模板文件
        :param module_path: 模板文件路径
        :param data_path: .fre文件路径
        :param output_path: 报告输出位置
        :return:
        """
        new_doc_path = self.module_path
        # 将doc文件转为docx文件，读取，
        # 传入3个参数、模板路径，数据路径，输出路径

            # new_doc_path = self.to_docx(new_doc_path)

        dom = docx.Document(new_doc_path)

        if not dom:
            print('打开word文档出错')
            return False
        output_data = open(os.path.dirname(self.output_path)+'\\'+'应力分析专用数据.tmp', 'w+', encoding='utf8')
        if self.data_path:
            self.deal_fre_table(dom)   # 处理fre数据功能函数
        self.deal_image(dom)  # 处理图片函数功能
        if self.max_path:
            print('max1223232')
            self.deal_max_table(dom, output_data)   # 处理。max文件数据
            print('max1')
        if self.prr_path:
            self.deal_prr_table(dom, output_data)   # 处理。prr文件数据
            print('max2')
        if self.ppo_path:
            self.deal_appendix_1(dom, output_data)  # 添加附录1
            print('max3')
        if self.prc_path:
            self.deal_appendix_2(dom, output_data)  # 添加附录2
            print('max4')
        output_data.close()  # 关闭文件
        self.list_equipment_nodes.clear()
        if os.path.exists(self.output_path):
            os.remove(self.output_path)
        try:    # 捕获因word未能打开，或正常读写的异常*/+
            dom.save(self.output_path)
            return True
        except Exception as e:
            return False

    def deal_fre_table(self, dom):
        tables = dom.tables
        table = tables[2]
        table2_4 = tables[4]
        table3_1_3 = tables[7]
        # 获取2.1节表格
        #  管道号    材料    外径(mm)    壁厚(mm)   保温线重(kg/m)    规范等级    抗震类别
        # 数据来源：.fre文件          OD为外径，WT为壁厚，MA为保温线重
        # MATERIAL TABLE HEADER
        # CROS CD=1 OD=76 WT=3.5 MA=6.26 SO=1 ST=1 KL=1
        # CROS CD=2 OD=89 WT=3.5 MA=7.38 SO=1 ST=1 KL=1
        more = 0
        more1 = 0
        is_more = False
        is_more1 = False
        with open(self.data_path, 'r', errors='ignore') as f:
            while True:
                line = f.readline()
                if line.find('* MATERIAL:') == 0:
                    material = line.split(':')[1].split('(')[0].strip()
                elif 'OD=' in line and 'WT=' in line and 'MA=' in line:

                    od = self.analysis_data(line, 'OD=')
                    wt = self.analysis_data(line, 'WT=')
                    ma = self.analysis_data(line, 'MA=')

                    if len(table.rows) < more + 2:
                        table.add_row()
                    for i in range(1, len(table.rows)):
                        if table.cell(i, 2).text.strip() == od and table.cell(i, 3).text.strip() == wt and \
                                table.cell(i, 4).text.strip() == ma:
                            is_more = True
                            break
                    if is_more:
                        is_more = False
                        continue
                    table.cell(more + 1, 2).text = od
                    table.cell(more + 1, 3).text = wt
                    table.cell(more + 1, 4).text = ma
                    more += 1
                elif 'VALV' in line and 'PT' in line and 'MA' in line:
                    pt = self.analysis_data(line, 'PT=')
                    ma = float(self.analysis_data(line, 'MA='))

                    dx = 0
                    dy = 0
                    dz = 0
                    if 'DX=' in line:
                        dx = float(self.analysis_data(line, 'DX='))
                    if 'DY=' in line:
                        dy = float(self.analysis_data(line, 'DY='))
                    if 'DZ=' in line:
                        dz = float(self.analysis_data(line, 'DZ='))

                    length = str(round((dx**2 + dy**2 + dz**2)**0.5, 3))
                    weight = str(round(ma*1000, 3))
                    if len(table2_4.rows) < more1 + 2:
                        table2_4.add_row()
                    for i in range(1, len(table2_4.rows)):
                        if table2_4.cell(i, 2).text.strip() == pt and table2_4.cell(i, 3).text.strip() == length and \
                                table2_4.cell(i, 4).text.strip() == weight:
                            is_more1 = True
                            break
                    if is_more1:
                        is_more1 = False
                        continue
                    table2_4.cell(more1 + 1, 2).text = pt
                    table2_4.cell(more1 + 1, 3).text = length
                    table2_4.cell(more1 + 1, 4).text = weight
                    table2_4.cell(more1 + 1, 0).text = str(more1+1)
                    more1 += 1

                elif re.search('NOZZ', line, re.IGNORECASE):
                    temp = "ACTION OF PIPING ON RESTRAINT(S) AT POINT NO.  " + line.split()[1].split('=')[1]. \
                        rjust(3, ' ')
                    self.list_equipment_nodes.append(temp)

                elif 'SPEC' in line and line.startswith('SPEC'):
                    all_spec = '    ' + line
                    while True:
                        line = f.readline()

                        if line.startswith('*'):
                            break
                        if line != '\n':

                            all_spec += line

                    for par in dom.paragraphs:
                        if '1）冲击载荷' in par.text:
                            par.runs[-1].add_break()
                            run = par.add_run(all_spec)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            par.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT

                elif 'LCAS' in line and 'CA=71' in line:
                    gx, gy, gz = '', '', ''
                    if self.analysis_data(line, 'GX='):
                        gx = 'GX=' + self.analysis_data(line, 'GX=')
                    if self.analysis_data(line, 'GY='):
                        gy = 'GY=' + self.analysis_data(line, 'GY=')
                    if self.analysis_data(line, 'GZ='):
                        gz = 'GZ=' + self.analysis_data(line, 'GZ=')
                    rolling_line = '    ' + '横摇：' + ' ' + gx + ' ' + gy + ' ' + gz + '\n'
                    for par in dom.paragraphs:
                        if '2）摇摆载荷' in par.text:
                            par.runs[-1].add_break()
                            run = par.add_run(rolling_line)
                            run.font.name = 'SimSun'
                            run.font.size = Pt(12)
                            par.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT

                elif 'LCAS' in line and 'CA=73' in line:
                    gx, gy, gz = '', '', ''
                    if self.analysis_data(line, 'GX='):
                        gx = 'GX=' + self.analysis_data(line, 'GX=')
                    if self.analysis_data(line, 'GY='):
                        gy = 'GY=' + self.analysis_data(line, 'GY=')
                    if self.analysis_data(line, 'GZ='):
                        gz = 'GZ=' + self.analysis_data(line, 'GZ=')
                    picth_line = '    ' + '纵摇：' + ' ' + gx + ' ' + gy + ' ' + gz + '\n'
                    for par in dom.paragraphs:
                        if '2）摇摆载荷' in par.text:
                            par.runs[-1].add_break()
                            run = par.add_run(picth_line)
                            run.font.name = 'SimSun'
                            run.font.size = Pt(12)
                            par.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT

                elif 'DESN' in line:
                    te, pr = '', ''
                    if self.analysis_data(line, 'TE='):
                        te = self.analysis_data(line, 'TE=')
                    if self.analysis_data(line, 'PR='):
                        pr = self.analysis_data(line, 'PR=')
                    table3_1_3.cell(2, 1).text = te
                    table3_1_3.cell(2, 2).text = pr

                elif 'OPER' in line and 'CA=20' in line:
                    te, pr = '', ''
                    if self.analysis_data(line, 'TE='):
                        te = self.analysis_data(line, 'TE=')
                    if self.analysis_data(line, 'PR='):
                        pr = self.analysis_data(line, 'PR=')
                    table3_1_3.cell(2, 3).text = te
                    table3_1_3.cell(2, 4).text = pr

                elif 'OPER' in line and 'CA=30' in line:
                    te, pr = '', ''
                    if self.analysis_data(line, 'TE='):
                        te = self.analysis_data(line, 'TE=')
                    if self.analysis_data(line, 'PR='):
                        pr = self.analysis_data(line, 'PR=')
                    table3_1_3.cell(2, 5).text = te
                    table3_1_3.cell(2, 6).text = pr

                elif 'OPER' in line and 'CA=40' in line:
                    te, pr = '', ''
                    if self.analysis_data(line, 'TE='):
                        te = self.analysis_data(line, 'TE=')
                    if self.analysis_data(line, 'PR='):
                        pr = self.analysis_data(line, 'PR=')
                    table3_1_3.cell(2, 7).text = te
                    table3_1_3.cell(2, 8).text = pr

                elif 'OPER' in line and 'CA=50' in line:
                    te, pr = '', ''
                    if self.analysis_data(line, 'TE='):
                        te = self.analysis_data(line, 'TE=')
                    if self.analysis_data(line, 'PR='):
                        pr = self.analysis_data(line, 'PR=')
                    table3_1_3.cell(2, 9).text = te
                    table3_1_3.cell(2, 10).text = pr

                if not line:
                    break
            for i in range(1, len(table.rows)):
                table.cell(i, 1).text = material

    @staticmethod
    def analysis_data(var_str, goal):
        """
        解析字符串
        :param var_str:
        :param goal:
        :return:
        """
        datas = var_str.split()
        for d in datas:
            if goal in d:
                data = d[len(goal):]
                return data

    def deal_image(self, doc):
        """
        插入图片
        图片来源：工作文件夹下文件名为‘model’得图片文件
        """
        pic_path = ''  # 图片路径
        for root, dirs, files in os.walk(os.path.dirname(self.data_path)):
            for file in files:
                try:
                    if os.path.splitext(file)[1] == '.jpg' and os.path.splitext(file)[0] == 'model':
                        pic = os.path.join(root, file)
                        pic_path = pic
                except Exception:
                    raise Exception('没有图片文件')
        for par in doc.paragraphs:
            if '本报告中涉及的管道模型如下图所示：' in par.text:
                try:
                    par.runs[-1].add_break()
                    par.runs[-1].add_picture(pic_path, width=Inches(4.8))
                    par.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
                except Exception:
                    raise Exception('图片无法插入')
                finally:
                    return

    def deal_max_table(self, doc, output_data):

        tables = doc.tables
        # 获取5.1节表格
        table5_1_1 = tables[19]
        # 模型中节点号	冲击（OBS）摇摆应力(MPa)	公式9-AB (MPa)	应力比
        # 数据来源：.max文件
        table5_1_2 = tables[20]
        # 模型中节点号    冲击（SSS）和摇摆应力(MPa) 公式9 - D(MPa) 应力比
        table5_2_1 = tables[23]

        str_data = StringIO()
        is_write_info = False
        value22_1 = 0
        value22_2 = 1
        value23_1 = 0
        value23_2 = 1
        value_temp_1 = 0
        value_temp_2 = 0
        read_line = ""
        value5_1 = ''

        with open(self.max_path, 'r', encoding='utf8') as f:
            while True:
                line = f.readline().strip('\n')
                # print(line)
                if "CALCULATION NUMBER" in line:
                    is_write_info = True
                    str_data.truncate(0)
                    str_data.seek(0)
                if is_write_info:
                    str_data.write(line+'\n')
                    
                if "LOADING CASE NO. 108" in line:
                    print(108)
                    output_data.write("@O级准则-方程8"+"\n")
                    for i in range(13):
                        line = f.readline()
                        str_data.write(line)

                    line = f.readline()
                    str_data.write(line)

                    for i in range(9):
                        str_data.write(f.readline())
                    output_data.write(str_data.getvalue())
                    is_write_info = False

                    data_2, data_11, data_13, data_14 = self.analysis_line(line)
                    table5_2_1.cell(1, 2).text = data_2
                    value22_1 = float(data_11)+float(data_13)
                    table5_2_1.cell(1, 3).text = format(value22_1, '.2f')
                    value22_2 = float(data_14)
                    table5_2_1.cell(1, 4).text = format(value22_1/value22_2, '.2f')
                    table5_2_1.cell(1, 5).text = data_14
                    print(1088)

                if "LOADING CASE NO. 110" in line:
                    print(110)
                    output_data.write("@AB级准则 方程10"+'\n')
                    str_data.truncate()

                    for i in range(15):
                        line = f.readline()
                        str_data.write(line)

                    line = f.readline()
                    str_data.write(line)

                    for i in range(9):
                        str_data.write(f.readline())

                    output_data.write(str_data.getvalue())
                    is_write_info = False

                    data_2, data_11, data_12, data_13 = self.analysis_line2(line)
                    table5_2_1.cell(3, 2).text = data_2
                    value22_1 = float(data_11) + float(data_12)
                    table5_2_1.cell(3, 3).text = format(value22_1, '.2f')
                    value22_2 = float(data_13)
                    table5_2_1.cell(3, 4).text = format(value22_1 / value22_2, '.2f')
                    table5_2_1.cell(3, 5).text = data_13
                    print(110110)

                if "LOADING CASE NO. 111" in line:
                    print(111)
                    output_data.write("@AB级准则 方程11"+'\n')

                    for i in range(13):
                        line = f.readline()
                        str_data.write(line)

                    line = f.readline()
                    str_data.write(line)

                    for i in range(9):
                        str_data.write(f.readline())

                    output_data.write(str_data.getvalue())
                    is_write_info = False

                    data_2, data_11, data_12, data_13 = self.analysis_line2(line)

                    table5_2_1.cell(4, 2).text = data_2
                    value22_1 = float(data_11) + float(data_12)
                    table5_2_1.cell(4, 3).text = format(value22_1, '.2f')
                    value22_2 = float(data_13)
                    table5_2_1.cell(4, 4).text = format(value22_1 / value22_2, '.2f')
                    table5_2_1.cell(4, 5).text = data_13
                    print(111111)

                if "LOADING CASE NO.  209 " in line:
                    print(209)
                    output_data.write("@单独由冲击（OBS）和摇摆作用下的管道最大应力比"+'\n')

                    for i in range(11):
                        line = f.readline()
                        str_data.write(line)

                    line = f.readline()
                    str_data.write(line)

                    for i in range(9):
                        str_data.write(f.readline())

                    output_data.write(str_data.getvalue())
                    is_write_info = False

                    data_2, data_11, data_13, data_14 = self.analysis_line(line)
                    table5_1_1.cell(1, 0).text = data_2
                    value22_1 = float(data_11) + float(data_13)
                    value5_1 = value22_1
                    table5_1_1.cell(1, 1).text = format(value22_1, '.2f')
                    print(209209)

                if "LOADING CASE NO. 109" in line:
                    print(109)
                    output_data.write("@单独由冲击（OBS）和摇摆作用下的管道最大应力比"+'\n')

                    for i in range(13):
                        line = f.readline()
                        str_data.write(line)

                    line = f.readline()
                    str_data.write(line)

                    for i in range(9):
                        str_data.write(f.readline())

                    output_data.write(str_data.getvalue())
                    output_data.write("@AB级准则-方程9"+'\n')
                    output_data.write(str_data.getvalue())
                    is_write_info = False

                    data_2, data_11, data_13, data_14 = self.analysis_line(line)
                    value22_2 = float(data_11) + float(data_13)
                    table5_1_1.cell(1, 2).text = format(value22_2, '.2f')

                    table5_2_1.cell(2, 2).text = data_2
                    value_temp_1 = value22_2
                    table5_2_1.cell(2, 3).text = format(value_temp_1, '.2f')
                    value_temp_2 = float(data_14)
                    table5_2_1.cell(2, 4).text = format(value_temp_1/value_temp_2, '.2f')
                    table5_2_1.cell(2, 5).text = data_14
                    print(109109)

                if "LOADING CASE NO.  709 " in line:
                    print(709)
                    output_data.write("@单独由冲击（SSS）和摇摆作用下的管道最大应力比" + '\n')

                    for i in range(11):
                        line = f.readline()
                        str_data.write(line)

                    line = f.readline()
                    str_data.write(line)

                    for i in range(9):
                        str_data.write(f.readline())

                    output_data.write(str_data.getvalue())
                    is_write_info = False

                    data_2, data_11, data_13, data_14 = self.analysis_line(line)
                    table5_1_2.cell(1, 0).text = data_2
                    value23_1 = float(data_11) + float(data_13)
                    table5_1_2.cell(1, 1).text = format(value23_1, '.2f')
                    print(709709)

                if "LOADING CASE NO. 140" in line:
                    print(140)
                    output_data.write("@C级准则 方程9"+'\n')

                    for i in range(13):
                        line = f.readline()
                        str_data.write(line)

                    line = f.readline()
                    str_data.write(line)

                    for i in range(9):
                        str_data.write(f.readline())

                    output_data.write(str_data.getvalue())
                    is_write_info = False

                    data_2, data_11, data_13, data_14 = self.analysis_line(line)
                    table5_2_1.cell(5, 2).text = data_2
                    value_temp_1 = float(data_11) + float(data_13)
                    table5_2_1.cell(5, 3).text = format(value_temp_1, '.2f')
                    value_temp_2 = float(data_14)
                    table5_2_1.cell(5, 4).text = format(value_temp_1 / value_temp_2, '.2f')
                    table5_2_1.cell(5, 5).text = data_14
                    print(140140)

                if "LOADING CASE NO. 150" in line:
                    print(150)
                    output_data.write("@D级准则 方程9"+'\n')

                    for i in range(13):
                        line = f.readline()
                        str_data.write(line)

                    line = f.readline()
                    str_data.write(line)

                    for i in range(9):
                        str_data.write(f.readline())

                    output_data.write(str_data.getvalue())
                    output_data.write("@单独由冲击（SSS）和摇摆作用下的管道最大应力比"+'\n')
                    output_data.write(str_data.getvalue())
                    is_write_info = False

                    data_2, data_11, data_13, data_14 = self.analysis_line(line)
                    value23_2 = float(data_11) + float(data_13)
                    table5_1_2.cell(1, 2).text = format(value23_2, '.2f')

                    table5_2_1.cell(6, 2).text = data_2
                    value_temp_1 = value23_2
                    table5_2_1.cell(6, 3).text = format(value_temp_1, '.2f')
                    value_temp_2 = float(data_14)
                    table5_2_1.cell(6, 4).text = format(value_temp_1/value_temp_2, '.2f')
                    table5_2_1.cell(6, 5).text = data_14
                    print(150150)

                if "LOADING CASE NO.  16" in line:
                    print(16)
                    output_data.write("@水压试验"+'\n')

                    for i in range(11):
                        line = f.readline()
                        str_data.write(line)

                    line = f.readline()
                    str_data.write(line)

                    for i in range(9):
                        str_data.write(f.readline())

                    output_data.write(str_data.getvalue())
                    is_write_info = False

                    # data = line.split()
                    data_2, data_11, data_13, data_14 = self.analysis_line(line)

                    table5_2_1.cell(7, 2).text = data_2
                    value_temp_1 = float(data_11) + float(data_13)
                    table5_2_1.cell(7, 3).text = format(value_temp_1, '.2f')
                    value_temp_2 = float(data_14)
                    table5_2_1.cell(7, 4).text = format(value_temp_1/value_temp_2, '.2f')
                    table5_2_1.cell(7, 5).text = data_14
                    print(1616)

                if not line:
                    break
    
            print(value5_1, value22_2, value23_1, value23_2)
            if value5_1 and value22_2:
                table5_1_1.cell(1, 3).text = format(value5_1 / value22_2, '.2f')
            if value23_1 and value23_2:
                table5_1_2.cell(1, 3).text = format(value23_1 / value23_2, '.2f')

    def deal_prr_table(self, doc, output_data):
        # 获取5.1节表格  共2个表格需要处理
        # 模型中节点	X	Y	Z
        # 数据来源：.prr文件

        tables = doc.tables
        table_1 = tables[21]  # 获取word中表格位置5_1_3
        table_2 = tables[22]  # 获取word中表格位置5_1_4

        str_data = StringIO()
        str_data1 = StringIO()
        str_data2 = StringIO()
        temp_data = StringIO()
        temp_data2 = StringIO()
        is_right = False
        first_1 = False
        first_2 = False
        first_11 = False
        first_22 = False
        j_30 = 0
        j_40 = 0
        with open(self.prr_path, 'r') as f:
            while True:
                line = f.readline()
                if "CALCULATION NUMBER" in line:
                    str_data.truncate(0)
                    str_data.seek(0)
                    str_data.write(line + '\n')
                    str_data.write(f.readline())
                    str_data.write(f.readline())

                if "LOADING CASE NO.   2" in line:
                    if not first_1:
                        first_1 = True
                        str_data1.write("@OBS冲击摇摆计算得到的阀门最大加速度"+"\n")
                        str_data1.write(str_data.getvalue())
                        str_data.truncate(0)
                    temp_data.write(line)

                    for i in range(8):
                        line = f.readline()
                        if first_1:
                            temp_data.write(line)
                        if "GLOBAL BOUNDS" in line:
                            is_right = True

                    if not is_right:
                        temp_data.truncate(0)
                        temp_data.seek(0)
                        continue

                    if not first_11:
                        is_right = False
                        first_11 = True
                        str_data1.write(temp_data.getvalue())

                    temp_data.truncate(0)

                    while True:
                        line = f.readline()
                        if line[0] != ' ':
                            break

                        if "VALVE" in line:
                            if len(table_1.rows) < j_30 + 2:
                                table_1.add_row()
                            str_data1.write(line)
                            table_1.cell(j_30 + 1, 0).text = line[3:9].strip()
                            table_1.cell(j_30 + 1, 1).text = line[72:81].strip()
                            table_1.cell(j_30 + 1, 2).text = line[81:90].strip()
                            table_1.cell(j_30 + 1, 3).text = line[90:99].strip()
                            j_30 += 1

                if "LOADING CASE NO.   9" in line:
                    if not first_2:
                        first_2 = True
                        str_data2.write("@SSS冲击摇摆计算得到的阀门最大加速度"+"\n")
                        str_data2.write(str_data.getvalue())
                        str_data.truncate(0)

                    temp_data2.write(line)

                    for i in range(8):
                        line = f.readline()
                        if first_2:
                            temp_data2.write(line)
                        if "GLOBAL BOUNDS" in line:
                            is_right = True

                    if not is_right:
                        temp_data2.truncate(0)
                        temp_data2.seek(0)
                        continue

                    if not first_22:
                        first_22 = True
                        is_right = False
                        str_data2.write(temp_data2.getvalue())

                    # temp_data2.truncate(0)

                    while True:
                        line = f.readline()
                        if line[0] != ' ':
                            break

                        if "--VALVE--" in line:
                            if len(table_2.rows) < j_40 + 2:
                                table_2.add_row()
                            str_data2.write(line)
                            table_2.cell(j_40 + 1, 0).text = line[3:9].strip()
                            table_2.cell(j_40 + 1, 1).text = line[72:81].strip()
                            table_2.cell(j_40 + 1, 2).text = line[81:90].strip()
                            table_2.cell(j_40 + 1, 3).text = line[90:99].strip()
                            j_40 += 1
                if not line:
                    break

        output_data.write(str_data1.getvalue())
        output_data.write(str_data2.getvalue())

    def deal_appendix_1(self, dom, output_data):

        par1 = dom.add_paragraph('附录1 支撑载荷')

        is_first = False  # 记录开始位置
        is_second = False  # 记录是否为需要的数据
        is_switch = False  # 是否是设备接管节点号
        str_data = StringIO()
        str_temp = ""
        index = 0

        with open(self.ppo_path, 'r') as f:
            while True:
                line = f.readline()    # 读取行内容
                if line.startswith('='):
                    is_first = True
                    index = 0
                    is_second = False
                    str_data.truncate(0)
                    str_data.seek(0)
                    str_data.write(line)

                elif is_first:
                    if line.startswith('-'):
                        index += 1
                    str_data.write(line)
                    if "ACTION OF PIPING ON RESTRAINT(S) AT POINT NO.  " in line:
                        str_temp = line
                        is_second = True
                    if 2 == index and is_second:
                        is_first = False
                        for i in range(len(self.list_equipment_nodes)):
                            if self.list_equipment_nodes[i] in str_temp:
                                is_switch = True
                                # output_data.write("@设备接管+节点" + str_temp[47:51].strip())
                                # output_data.write(str_data.getvalue())
                                # str_data1.write(str_data.getvalue())
                                break
                        if not is_switch:
                            output_data.write("@支撑+节点" + str_temp[47: 51].strip()+'\n')
                            output_data.write(str_data.getvalue())
                            par1.paragraph_format.keep_together = True
                            par1.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED
                            par1.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            par1.runs[-1].add_break()
                            run = par1.add_run(str_data.getvalue())
                            run.font.size = Pt(7)
                            run.font.name = 'Times New Roman'
                            par1.runs[-1].add_break()

                    is_switch = False
                if not line:
                    break
        is_switch = False
        is_first = False
        is_second = False
        str_data3 = StringIO()
        str_temp1 = ''
        par2 = dom.add_paragraph('附录2 接管载荷')

        with open(self.ppo_path, 'r') as f:
            while True:
                line = f.readline().strip('\n')
                if line.startswith("="):
                    is_first = True
                    index = 0
                    is_second = False
                    str_data3.truncate(0)
                    str_data3.seek(0)
                    str_data3.write(line+'\n')
                elif is_first:
                    if line.startswith('-'):
                        index += 1
                    str_data3.write(line + '\n')
                    if "ACTION OF PIPING ON RESTRAINT(S) AT POINT NO.  " in line:
                        str_temp1 = line
                        is_second = True
                    if 2 == index and is_second:
                        is_first = False
                        for i in range(len(self.list_equipment_nodes)):
                            if self.list_equipment_nodes[i] in str_temp1:
                                is_switch = True
                                output_data.write("@设备接管+节点" + str_temp1[47:51].strip() + '\n')
                                output_data.write(str_data3.getvalue())

                                par2.runs[-1].add_break()
                                par2.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED
                                par2.line_spacing_rule = WD_LINE_SPACING.SINGLE
                                run = par2.add_run(str_data3.getvalue())
                                run.font.size = Pt(7)
                                run.font.name = 'Times New Roman'
                                par2.runs[-1].add_break()

                                str_data3.truncate(0)
                                str_data3.seek(0)
                                break

                if not line:
                    break

    def deal_appendix_2(self, dom, output_data):

        par3 = dom.add_paragraph('附录3 位移输出')

        index = 0
        num = 0
        is_first = False  # 记录开始位置
        is_second = False  # 记录是否为需要的数据
        is_third = False  # 是否是设备接管节点号
        is_many_data = False
        str_data = StringIO()
        str_data1 = StringIO()

        with open(self.prc_path, 'r', encoding='utf8') as f:
            while True:
                line = f.readline()
                if "CALCULATION NUMBER" in line:  # 寻找标志数据
                    is_first = True
                    index = 0
                    str_data.truncate(0)
                    str_data.seek(0)

                    if not is_many_data:
                        str_data.write(line)
                elif is_first:
                    if not line.startswith(' '):
                        index += 1
                    if "LOADING CASE NO. 150" in line:
                        is_second = True
                    if "GLOBAL ---DISPLACEMENTS (MM)---    --ROTATIONS (RAD*1000)--       ------ACCEL (G-S)------" in line:
                        is_third = True
                    if 1 == index & is_second & is_third:
                        is_first = False
                        is_second = False
                        is_third = False
                        is_many_data = True
                        num = 0
                        str_data1.write(str_data.getvalue())
                    if is_many_data:
                        num += 1
                        if num > 11:
                            str_data.write(line)
                    else:
                        str_data.write(line)
                if not line:
                    break
        output_data.write("@位移输出" + '\n')
        if str_data1:
            par3.runs[-1].add_break()
            run = par3.add_run(str_data1.getvalue())
            run.font.size = Pt(7)
            run.font.name = 'Times New Roman'
            par3.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED
            output_data.write(str_data1.getvalue())

        else:
            output_data.write("empty")


# if __name__ == '__main__':
#     # f1 = r'D:\peps\peps资料\PEPS二次开发参考资料\PEPS输入输出文件\cztsg.fre'
#     # f2 = r'D:\peps\peps资料\PEPS二次开发参考资料\力学分析报告模板.docx'
#     # f3 = r'D:\peps\peps资料\PEPS二次开发参考资料\TEST_WORD.docx'
#     #
#     module_p = r'D:\peps\CASE\力学模板.docx'
#     fre_p = r'D:\peps\CASE\case-PRHRS-WATER-1.fre'
#     new_report = r'D:\peps\CASE\力学报告2021.docx'
#
#     obj = ExportReport(module_p, fre_p, new_report)
# #     # obj = ExportReport(f2, f1, f3)
